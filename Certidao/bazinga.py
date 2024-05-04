import docx
import os

def read_data_from_txt(file_path):
    """ Lê os dados de um arquivo TXT e retorna um dicionário com as chaves e valores. """
    if not os.path.exists(file_path):
        print(f"Erro: Arquivo não encontrado {file_path}")
        return {}
    data = {}
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            if ': ' in line:
                parts = line.strip().split(': ', 1)
                if len(parts) == 2:  # Garante que a linha contém uma chave e um valor
                    key, value = parts
                    data[key] = value
                else:
                    print(f"Erro: Linha mal formatada, não pode ser dividida em chave e valor -> {line.strip()}")
    return data

def find_and_replace(doc, key, value):
    """ Procura e substitui a primeira ocorrência do texto chave por valor, em parágrafos e tabelas. """
    found = False
    for paragraph in doc.paragraphs:
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)
            found = True
            break

    if not found:  # Procura nas tabelas apenas se não foi encontrado nos parágrafos
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                            found = True
                            return True  # Retorna assim que fizer a primeira substituição
    return found

def replace_in_docx(template_path, output_path, data):
    """ Lê um template DOCX, substitui texto conforme especificado e salva o resultado. """
    if not os.path.exists(template_path):
        print(f"Erro: Arquivo de template não encontrado {template_path}")
        return

    doc = docx.Document(template_path)
    replaced = False

    # Mapeamento de chaves de dados para marcadores no documento Word
    mappings = {
        'NAME': '#NAME',
        'SOCIAL SECURITY CARD': '#SOCIALSECURITY',
        'REGISTRATION': '#REGISTRATION',
        'DATE AND TIME OF BIRTH': '#DATETIMEBIRTH',
        'CITY AND STATE OF BIRTH': '#CITYSTATEBIRTH',
        'PLACE OF BIRTH, CITY AND STATE': '#PLACEBIRTHCITYSTATE',
        'CITY AND STATE OF REGISTRATION': '#CITYSTATEREGISTRATION',
        'GENDER': '#GENDER',
        'MOTHER': '#MOTHER',
        'MATERNAL GRANDPARENTS': '#MATERNALGRADPARENTS',
        'TWINS': '#TWINS',
        'NAME AND REGISTRATION OF TWINS': '#NAMEREGISTRATIONTWINS',
        'DATE OF REGISTRATION': '#DATETWINS',
        'REGISTRATION NUMBER OF ALIVE BIRTH': '#REGISTRATIONNUMBERALIVEBIRTH',
        'NOTES AMENDMENTS TO ADD': '#AMENDMENT',
        'FATHER': '#FATHER',
        'PATERNAL GRANDPARENTS': '#PATERNALGRANDPARENTS',
        'DOCUMENT TYPE': '#DOCUMENTTYPE',
        'NUMBER': '#NUMBER',
        'ISSUANCE DATE': '#ISSUANCEDATE',
        'REGISTRATION NOTES': '#REGISTRATIONNOTES',
        'OBSERVATIONS': '#OBSERVATIONS'
    }


    for key, value in data.items():
        if key in mappings and find_and_replace(doc, mappings[key], value):
            replaced = True

    if replaced:
        doc.save(output_path)
        print("Arquivo DOCX processado com sucesso e salvo em:", output_path)
    else:
        print("Nenhuma substituição foi realizada.")

# Carrega dados e realiza substituição baseado em mapeamento
data = read_data_from_txt('input.txt')
if data:
    replace_in_docx('BCMODELO.docx', 'output.docx', data)
else:
    print("Nenhuma dado para processar.")

input("Pressione Enter para finalizar...")
