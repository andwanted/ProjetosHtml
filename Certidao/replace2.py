import docx
import os

def read_data_from_txt(file_path):
    """ Lê os dados de um arquivo TXT e retorna um dicionário com as chaves e valores. """
    if not os.path.exists(file_path):
        print(f"Erro: Arquivo não encontrado {file_path}")
        return {}
    data = {}
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            if ': ' in line:
                key, value = line.strip().split(': ', 1)
                data[key] = value
    return data

def replace_text_while_keeping_formatting(obj, key, value):
    """ Substitui texto em um objeto Document, mantendo a formatação original. """
    found = False
    if hasattr(obj, 'paragraphs'):
        for paragraph in obj.paragraphs:
            for run in paragraph.runs:
                if key in run.text:
                    print(f"Encontrado '{key}' em '{run.text}'. Substituindo por '{value}'")
                    run.text = run.text.replace(key, value)
                    found = True
                else:
                    print(f"'{key}' não encontrado em '{run.text}'")
    return found


def replace_in_docx(template_path, output_path, data):
    if not os.path.exists(template_path):
        print(f"Erro: Arquivo de template não encontrado {template_path}")
        return

    doc = docx.Document(template_path)
    replaced = False

    for p in doc.paragraphs:
        for key, value in data.items():
            if replace_text_while_keeping_formatting(p, key, value):
                replaced = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if replace_text_while_keeping_formatting(cell, key, value):
                        replaced = True

    if replaced:
        doc.save(output_path)
        print("Arquivo DOCX processado com sucesso e salvo em:", output_path)
    else:
        print("Nenhuma substituição foi realizada.")

# Mapeamento de chaves de dados para marcadores no documento Word
mappings = {
    'NAME': '#NAME',
    'SOCIAL SECURITY CARD': '#SOCIALSECURITY',
    'REGISTRATION': '#REGISTRATION',
    'DATE AND TIME OF BIRTH': '#DATETIMEBIRTH',
    'CITY AND STATE OF BIRTH': '#CITYSTATEBIRTH',
    'PLACE OF BIRTH, CITY AND STATE': '#PLACEBIRTHCITYSTATE',
    'CITY AND STATE OF REGISTRATION': '#PLACEBIRTHCITYSTATE',
    'GENDER': '#GENDER',
    'MOTHER': '#MOTHER',
    'MATERNAL GRANDPARENTS': '#MATERNALGRADPARENTS',
    'TWINS': '#TWINS',
    'NAME AND REGISTRATION OF TWINS': '#NAMEREGISTRATIONTWINS',
    'DATE OF REGISTRATION': '#DATEREGISTRIONTWINS',
    'REGISTRATION NUMBER OF ALIVE BIRTH': '#REGISTRATIONNUMBERALIVEBIRTH',
    'NOTES AMENDMENTS TO ADD': '#AMENDMENT',
    'ISSUANCE DATE': 'unknown',
    'FATHER': 'unknown',
    'PATERNAL GRANDPARENTS': 'unknown',
    'REGISTRATION NOTES': 'unknown',
    'DOCUMENT TYPE': 'unknown',
    'NUMBER': 'unknown',
    'OBSERVATIONS': 'unknown'
}

data = read_data_from_txt('input.txt')
if data:
    # Aplica o mapeamento
    mapped_data = {f'{mappings[key]}': value for key, value in data.items() if key in mappings}
    replace_in_docx('BCMODELO.docx', 'output.docx', mapped_data)
else:
    print("Nenhum dado para processar.")

input("Pressione Enter para finalizar...")
